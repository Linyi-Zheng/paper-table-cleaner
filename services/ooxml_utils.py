from __future__ import annotations

import re
import zipfile
from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Length, Pt
from lxml import etree

ALIGNMENT_BY_NAME = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "DISTRIBUTE": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}


def length_to_pt(value: Length | int | float | None, *, max_pt: float = 1000.0) -> float | None:
    if value is None:
        return None
    if hasattr(value, "pt"):
        pt_value = float(value.pt)
    else:
        raw_value = float(value)
        # python-docx lengths are EMUs. Some documents surface raw EMU integers.
        pt_value = raw_value / 12700 if abs(raw_value) > max_pt else raw_value
    if abs(pt_value) > max_pt:
        return None
    return round(pt_value, 2)


def pt_or_none(value: float | int | Length | None, *, max_pt: float = 1000.0) -> Length | None:
    if value is None:
        return None
    pt_value = length_to_pt(value, max_pt=max_pt)
    if pt_value is None:
        return None
    return Pt(float(pt_value))


def font_size_to_pt(value: Length | int | float | None) -> float | None:
    return length_to_pt(value, max_pt=200.0)


def normalize_line_spacing(value: object | None) -> float | None:
    if value is None:
        return None
    if hasattr(value, "pt"):
        return length_to_pt(value, max_pt=200.0)
    if isinstance(value, (int, float)):
        raw_value = float(value)
        if 0 < raw_value <= 10:
            return round(raw_value, 2)
        converted = raw_value / 12700
        if 0 < converted <= 200:
            return round(converted, 2)
    return None


def apply_line_spacing(paragraph_format: object, value: float | int | None) -> None:
    if value is None:
        return
    spacing = float(value)
    if spacing <= 0:
        return
    if spacing <= 10:
        paragraph_format.line_spacing = spacing
        return
    if spacing <= 200:
        paragraph_format.line_spacing = Pt(spacing)


def alignment_to_name(value: object | None) -> str | None:
    if value is None:
        return None
    name = getattr(value, "name", None)
    if name:
        return str(name)
    return str(value)


def name_to_alignment(value: str | None) -> object | None:
    if not value:
        return None
    return ALIGNMENT_BY_NAME.get(value.upper())


def normalize_text(text: str) -> str:
    normalized = text.replace("\u3000", " ").replace("□", " ")
    return re.sub(r"\s+", " ", normalized).strip()


def _set_rfonts(r_fonts: object, east_asia_font: str | None, latin_font: str | None) -> None:
    if latin_font:
        for key in ("w:ascii", "w:hAnsi", "w:cs"):
            r_fonts.set(qn(key), latin_font)
    if east_asia_font:
        r_fonts.set(qn("w:eastAsia"), east_asia_font)


def set_run_font(run: object, font_name: str | None, latin_font_name: str | None = None) -> None:
    if not font_name and not latin_font_name:
        return
    run.font.name = latin_font_name or font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    _set_rfonts(r_fonts, font_name, latin_font_name or font_name)


def set_style_font(style: object, font_name: str | None, latin_font_name: str | None = None) -> None:
    if not font_name and not latin_font_name:
        return
    style.font.name = latin_font_name or font_name
    element = getattr(style, "_element", None)
    if element is None:
        element = getattr(style, "element", None)
    if element is None:
        return
    r_pr = element.get_or_add_rPr() if hasattr(element, "get_or_add_rPr") else element.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        element.append(r_pr)
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    _set_rfonts(r_fonts, font_name, latin_font_name or font_name)


def clear_paragraph_runs(paragraph: object) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_container_plain_text(container: object, text: str | None) -> None:
    paragraphs = list(container.paragraphs)
    if not paragraphs:
        paragraph = container.add_paragraph()
        paragraphs = [paragraph]
    for paragraph in paragraphs:
        clear_paragraph_runs(paragraph)
    if text:
        paragraphs[0].add_run(text)


def copy_header_footer_plain_text(
    target_section: object,
    header_text: str | None,
    footer_text: str | None,
) -> bool:
    try:
        set_container_plain_text(target_section.header, header_text)
        set_container_plain_text(target_section.footer, footer_text)
        return True
    except Exception:
        return False


def detect_page_number(container: object) -> bool:
    for node in container._element.iter():
        if node.tag == qn("w:instrText") and node.text:
            if "PAGE" in node.text.upper():
                return True
    text = "\n".join(paragraph.text for paragraph in container.paragraphs)
    return bool(re.search(r"\bPAGE\b|第\s*\d+\s*页|\d+\s*/\s*\d+", text, re.I))


def add_page_number_field(paragraph: object) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def insert_paragraph_after(paragraph: object, text: str = "", style: str | None = None):
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    if style:
        try:
            new_paragraph.style = style
        except Exception:
            pass
    return new_paragraph


def insert_paragraph_before(paragraph: object, text: str = "", style: str | None = None):
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    if style:
        try:
            new_paragraph.style = style
        except Exception:
            pass
    return new_paragraph


def has_toc_field(document: object) -> bool:
    for node in document.element.iter():
        if node.tag == qn("w:instrText") and node.text:
            if "TOC" in node.text.upper():
                return True
    return False


def add_toc_field(paragraph: object) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    instr_run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    paragraph.add_run("请在 Word 中右键更新域以生成目录。")

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def set_update_fields_on_open(document: object) -> None:
    settings = document.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def clone_paragraph_properties(source: object, target: object) -> None:
    source_p_pr = source._p.pPr
    if source_p_pr is None:
        return
    target_p_pr = target._p.get_or_add_pPr()
    target_p_pr.clear()
    for child in source_p_pr:
        target_p_pr.append(deepcopy(child))


def detect_complex_objects(docx_path: Path) -> list[str]:
    findings: list[str] = []
    if not zipfile.is_zipfile(docx_path):
        return ["文档不是有效的 OOXML zip 包。"]

    with zipfile.ZipFile(docx_path) as archive:
        names = archive.namelist()
        if any(name.startswith("word/embeddings/") for name in names):
            findings.append("检测到嵌入对象，已尽量保留但未自动修改。")
        if any(name.startswith("word/charts/") for name in names):
            findings.append("检测到图表对象，已尽量保留但未自动修改。")
        for name in names:
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            try:
                xml = archive.read(name)
                root = etree.fromstring(xml)
            except Exception:
                continue
            if root.xpath(".//*[local-name()='oMath' or local-name()='oMathPara']"):
                findings.append("检测到 Office 公式对象，已保留原内容。")
                break
    return list(dict.fromkeys(findings))


def first_non_empty_text(paragraphs: Iterable[object]) -> str | None:
    for paragraph in paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            return text
    return None


TABLE_BORDER_EDGES = ("top", "left", "bottom", "right", "insideH", "insideV")
CELL_BORDER_EDGES = ("top", "left", "bottom", "right", "insideH", "insideV")


def _get_or_add_child(parent: object, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _get_or_add_tbl_pr(table: object):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    return tbl_pr


def _set_border(parent: object, edge: str, *, val: str, size: int = 0, color: str = "auto") -> None:
    border = parent.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        parent.append(border)
    border.set(qn("w:val"), val)
    if val == "nil":
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
    else:
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def _get_or_add_table_borders(table: object):
    tbl_pr = _get_or_add_tbl_pr(table)
    return _get_or_add_child(tbl_pr, "w:tblBorders")


def set_table_width_to_window(table: object) -> None:
    """Make the table use the printable page width, like Word's AutoFit to Window."""

    table.autofit = True
    tbl_pr = _get_or_add_tbl_pr(table)
    tbl_w = _get_or_add_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    tbl_jc = _get_or_add_child(tbl_pr, "w:jc")
    tbl_jc.set(qn("w:val"), "center")

    tbl_ind = _get_or_add_child(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")

    # Explicitly disable Word's "Allow spacing between cells".
    # Positive spacing is the usual reason three-line table rules break at column boundaries.
    cell_spacing = _get_or_add_child(tbl_pr, "w:tblCellSpacing")
    cell_spacing.set(qn("w:type"), "dxa")
    cell_spacing.set(qn("w:w"), "0")

    layout = _get_or_add_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "autofit")

    column_count = max((len(row.cells) for row in table.rows), default=0)
    if column_count:
        cell_pct = str(int(5000 / column_count))
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = _get_or_add_child(tc_pr, "w:tcW")
                tc_w.set(qn("w:type"), "pct")
                tc_w.set(qn("w:w"), cell_pct)


def _get_or_add_cell_borders(cell: object):
    tc_pr = cell._tc.get_or_add_tcPr()
    return _get_or_add_child(tc_pr, "w:tcBorders")


def _remove_child(parent: object, tag: str) -> None:
    child = parent.find(qn(tag))
    if child is not None:
        parent.remove(child)


def _remove_table_borders(table: object) -> None:
    tbl_pr = _get_or_add_tbl_pr(table)
    _remove_child(tbl_pr, "w:tblBorders")


def _remove_cell_borders(cell: object) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    _remove_child(tc_pr, "w:tcBorders")


def _remove_all_cell_borders(table: object) -> None:
    for row in table.rows:
        for cell in row.cells:
            _remove_cell_borders(cell)


def clear_table_style_and_fills(table: object) -> None:
    """Remove table styles and cell shading that can create colored header rows."""

    tbl_pr = _get_or_add_tbl_pr(table)
    for tag in ("w:tblStyle", "w:tblLook", "w:shd", "w:tblCellSpacing"):
        _remove_child(tbl_pr, tag)
    for row in table.rows:
        tr_pr = row._tr.trPr
        if tr_pr is not None:
            _remove_child(tr_pr, "w:shd")
            _remove_child(tr_pr, "w:cnfStyle")
            _remove_child(tr_pr, "w:tblPrEx")
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            _remove_child(tc_pr, "w:shd")
            _remove_child(tc_pr, "w:cnfStyle")
            _remove_child(tc_pr, "w:tcBorders")
            for paragraph in cell.paragraphs:
                p_pr = paragraph._p.get_or_add_pPr()
                _remove_child(p_pr, "w:shd")
                for run in paragraph.runs:
                    r_pr = run._element.get_or_add_rPr()
                    _remove_child(r_pr, "w:highlight")
                    _remove_child(r_pr, "w:shd")


def clear_table_borders(table: object) -> None:
    _remove_table_borders(table)
    borders = _get_or_add_table_borders(table)
    for edge in TABLE_BORDER_EDGES:
        _set_border(borders, edge, val="nil")
    _remove_all_cell_borders(table)


def set_table_top_border(table: object, size: int = 12, color: str = "000000") -> None:
    borders = _get_or_add_table_borders(table)
    _set_border(borders, "top", val="single", size=size, color=color)


def set_table_bottom_border(table: object, size: int = 12, color: str = "000000") -> None:
    borders = _get_or_add_table_borders(table)
    _set_border(borders, "bottom", val="single", size=size, color=color)


def set_row_top_border(row: object, size: int = 12, color: str = "000000") -> None:
    for cell in row.cells:
        borders = _get_or_add_cell_borders(cell)
        _set_border(borders, "top", val="single", size=size, color=color)
        for edge in ("left", "right", "insideH", "insideV"):
            _set_border(borders, edge, val="nil")


def set_row_bottom_border(row: object, size: int = 8, color: str = "000000") -> None:
    for cell in row.cells:
        borders = _get_or_add_cell_borders(cell)
        _set_border(borders, "bottom", val="single", size=size, color=color)
        for edge in ("left", "right", "insideH", "insideV"):
            _set_border(borders, edge, val="nil")


def remove_vertical_borders(table: object) -> None:
    borders = _get_or_add_table_borders(table)
    for edge in ("left", "right", "insideV"):
        _set_border(borders, edge, val="nil")


def suppress_body_row_separator_borders(table: object) -> None:
    """Keep the header separator but suppress horizontal rules between body rows."""

    for row in list(table.rows)[2:]:
        for cell in row.cells:
            borders = _get_or_add_cell_borders(cell)
            _set_border(borders, "top", val="nil")
            for edge in ("left", "right", "insideH", "insideV"):
                _set_border(borders, edge, val="nil")


def apply_three_line_borders(
    table: object,
    *,
    top_size: int = 12,
    header_bottom_size: int = 8,
    bottom_size: int = 12,
    color: str = "000000",
) -> None:
    clear_table_style_and_fills(table)
    clear_table_borders(table)
    set_table_width_to_window(table)
    borders = _get_or_add_table_borders(table)
    for edge in TABLE_BORDER_EDGES:
        _set_border(borders, edge, val="nil")
    _remove_all_cell_borders(table)
    if table.rows:
        set_row_top_border(table.rows[0], top_size, color)
        set_row_bottom_border(table.rows[0], header_bottom_size, color)
        if len(table.rows) > 1:
            set_row_bottom_border(table.rows[-1], bottom_size, color)


def border_value(parent: object, edge: str) -> str | None:
    border = parent.find(qn(f"w:{edge}"))
    if border is None:
        return None
    return border.get(qn("w:val"))


def table_has_merged_cells(table: object) -> bool:
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.tcPr
            if tc_pr is None:
                continue
            if tc_pr.find(qn("w:gridSpan")) is not None or tc_pr.find(qn("w:vMerge")) is not None:
                return True
    return False
