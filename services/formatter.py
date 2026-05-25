from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.models import FormatRunResult, FormatSummary, ParagraphPrediction, StyleProfile, TemplateStyleProfile
from app.services.ooxml_utils import (
    add_page_number_field,
    apply_line_spacing,
    clear_paragraph_runs,
    copy_header_footer_plain_text,
    detect_complex_objects,
    font_size_to_pt,
    insert_paragraph_after,
    name_to_alignment,
    normalize_text,
    pt_or_none,
    set_run_font,
    set_style_font,
)
from app.services.figure_formatter import format_all_figures
from app.services.north_college_front_matter import (
    extract_front_matter_info,
    replace_north_college_front_matter,
)
from app.services.table_formatter import format_all_tables
from app.services.toc_builder import insert_or_mark_toc

BUILTIN_STYLE_BY_TYPE = {
    "body": "Normal",
    "heading_1": "Heading 1",
    "heading_2": "Heading 2",
    "heading_3": "Heading 3",
    "heading_4": "Heading 4",
}

PROFILE_BY_TYPE = {
    "body": "body",
    "heading_1": "heading_1",
    "heading_2": "heading_2",
    "heading_3": "heading_3",
    "heading_4": "heading_4",
    "contents_title": "heading_1",
    "references_title": "reference_title",
    "acknowledgements_title": "heading_1",
    "appendix_title": "heading_1",
}

SUMMARY_FIELD_BY_TYPE = {
    "body": "body_paragraphs_modified",
    "heading_1": "heading_1_modified",
    "heading_2": "heading_2_modified",
    "heading_3": "heading_3_modified",
    "heading_4": "heading_4_modified",
}

COVER_LABELS = {
    "姓名",
    "学号",
    "院系",
    "专业",
    "指导教师",
    "年月日",
    "学士学位论文",
    "学位论文原创性声明",
    "学位论文版权使用授权书",
}


def _front_matter_end_index(predictions: list[ParagraphPrediction]) -> int:
    for prediction in predictions:
        if prediction.predicted_type == "abstract_title":
            return prediction.index
    return 0


def _effective_prediction_type(prediction: ParagraphPrediction, protected_until: int) -> str | None:
    if prediction.index < protected_until:
        return None
    if prediction.predicted_type == "contents_entry":
        return None
    if prediction.predicted_type.startswith("heading") and prediction.confidence < 0.75:
        return "body"
    return prediction.predicted_type


def _compact_text(text: str) -> str:
    return re.sub(r"\s+", "", normalize_text(text))


def _remove_paragraph(paragraph: object) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
        paragraph._p = paragraph._element = None


def _is_cover_title_candidate(paragraph: object) -> bool:
    text = normalize_text(paragraph.text)
    compact = _compact_text(text)
    if not compact or compact in COVER_LABELS:
        return False
    if len(compact) < 10:
        return False
    if re.search(r"(声明|授权书|签名|年月日|论文作者|指导教师|保密|本人所提交)", text):
        return False
    if re.search(r"^\d{4}|^二[〇0零一二三四五六七八九十]+年", compact):
        return False
    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return True
    return any(run.text.strip() and (run.bold or run.underline) for run in paragraph.runs)


def _format_cover_heading(document: object, protected_until: int) -> bool:
    scan_until = protected_until if protected_until > 0 else min(30, len(document.paragraphs))
    heading_labels = {"学士学位论文", "本科毕业论文", "毕业论文", "学位论文"}
    for paragraph in document.paragraphs[:scan_until]:
        compact = _compact_text(paragraph.text)
        if compact not in heading_labels:
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.space_before = Pt(0)
        return True
    return False


def _set_paragraph_run_style(paragraph: object, profile: StyleProfile) -> None:
    for run in paragraph.runs:
        set_run_font(run, profile.font_name, profile.latin_font_name)
        font_size = font_size_to_pt(profile.font_size_pt)
        if font_size is not None:
            run.font.size = Pt(font_size)
        if profile.bold is not None:
            run.bold = profile.bold


def _format_cover_title(document: object, protected_until: int, template_profile: TemplateStyleProfile) -> bool:
    scan_until = protected_until if protected_until > 0 else min(30, len(document.paragraphs))
    candidates = [
        paragraph
        for paragraph in document.paragraphs[:scan_until]
        if _is_cover_title_candidate(paragraph)
    ]
    if not candidates:
        return False
    title = max(candidates, key=lambda paragraph: len(_compact_text(paragraph.text)))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.25 if template_profile.profile_name == "hebei_north_college_fixed" else 1.5
    title.paragraph_format.first_line_indent = None
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_together = True
    if template_profile.profile_name == "hebei_north_college_fixed":
        _set_paragraph_run_style(
            title,
            StyleProfile(
                font_name="黑体",
                latin_font_name="Times New Roman",
                font_size_pt=26,
                bold=True,
            ),
        )
    return True


def _compact_cover_spacing(document: object, protected_until: int) -> int:
    scan_until = protected_until if protected_until > 0 else min(30, len(document.paragraphs))
    paragraphs = list(document.paragraphs[:scan_until])
    heading = None
    title = None
    for paragraph in paragraphs:
        compact = _compact_text(paragraph.text)
        if compact in {"学士学位论文", "本科学位论文", "本科毕业论文", "毕业论文", "学位论文"}:
            heading = paragraph
            break
    candidates = [paragraph for paragraph in paragraphs if _is_cover_title_candidate(paragraph)]
    if candidates:
        title = max(candidates, key=lambda paragraph: len(_compact_text(paragraph.text)))
    if heading is None or title is None:
        return 0

    removed = 0
    seen_heading = False
    blanks_between: list[object] = []
    for paragraph in list(document.paragraphs[:scan_until]):
        if paragraph._element is heading._element:
            seen_heading = True
            continue
        if not seen_heading:
            continue
        if paragraph._element is title._element:
            break
        if not normalize_text(paragraph.text):
            blanks_between.append(paragraph)

    for paragraph in blanks_between[1:]:
        _remove_paragraph(paragraph)
        removed += 1
    return removed


def _format_cover_metadata(document: object, protected_until: int, template_profile: TemplateStyleProfile) -> int:
    if template_profile.profile_name != "hebei_north_college_fixed":
        return 0
    scan_until = protected_until if protected_until > 0 else min(30, len(document.paragraphs))
    formatted = 0
    for paragraph in document.paragraphs[:scan_until]:
        compact = _compact_text(paragraph.text)
        if not compact:
            continue
        if any(label in compact for label in ("姓名", "院系", "专业", "指导教师")):
            _set_paragraph_run_style(paragraph, StyleProfile(font_name="仿宋", latin_font_name="Times New Roman", font_size_pt=16))
            formatted += 1
        elif "学号" in compact:
            _set_paragraph_run_style(paragraph, StyleProfile(font_name="仿宋", latin_font_name="Times New Roman", font_size_pt=16))
            formatted += 1
        elif re.search(r"年.*月.*日", compact):
            _set_paragraph_run_style(paragraph, StyleProfile(font_name="宋体", latin_font_name="Times New Roman", font_size_pt=16))
            formatted += 1
    return formatted


def _format_declaration_signature_area(document: object, protected_until: int, template_profile: TemplateStyleProfile) -> int:
    if template_profile.profile_name != "hebei_north_college_fixed":
        return 0
    scan_until = protected_until if protected_until > 0 else min(80, len(document.paragraphs))
    signature_text = "论文作者（签名）：                  指导教师确认（签名）："
    date_text = "        年   月   日                    年   月   日"
    formatted = 0
    paragraphs = list(document.paragraphs[:scan_until])
    for index, paragraph in enumerate(paragraphs):
        compact = _compact_text(paragraph.text)
        if "论文作者" not in compact or "指导教师" not in compact or "签名" not in compact:
            continue

        _replace_paragraph_text(paragraph, signature_text)
        _apply_paragraph_profile(
            paragraph,
            StyleProfile(
                font_name="宋体",
                latin_font_name="Times New Roman",
                font_size_pt=12,
                bold=False,
                alignment="CENTER",
                first_line_indent=0,
                line_spacing=20,
                space_before=12,
                space_after=0,
            ),
        )
        formatted += 1

        next_paragraph = paragraphs[index + 1] if index + 1 < len(paragraphs) else None
        if next_paragraph is None or not re.search(r"年\s*月\s*日", normalize_text(next_paragraph.text)):
            next_paragraph = insert_paragraph_after(paragraph, date_text)
        else:
            _replace_paragraph_text(next_paragraph, date_text)
        _apply_paragraph_profile(
            next_paragraph,
            StyleProfile(
                font_name="宋体",
                latin_font_name="Times New Roman",
                font_size_pt=12,
                bold=False,
                alignment="CENTER",
                first_line_indent=0,
                line_spacing=20,
                space_before=0,
                space_after=0,
            ),
        )
        formatted += 1
        break
    return formatted


def _configure_document_style(document: object, style_name: str, profile: StyleProfile) -> None:
    try:
        style = document.styles[style_name]
    except KeyError:
        return
    set_style_font(style, profile.font_name, profile.latin_font_name)
    font_size = font_size_to_pt(profile.font_size_pt)
    if font_size:
        style.font.size = Pt(font_size)
    if profile.bold is not None:
        style.font.bold = profile.bold
    paragraph_format = style.paragraph_format
    if profile.line_spacing is not None:
        apply_line_spacing(paragraph_format, profile.line_spacing)
    if profile.first_line_indent is not None:
        paragraph_format.first_line_indent = pt_or_none(profile.first_line_indent)
    if profile.space_before is not None:
        paragraph_format.space_before = pt_or_none(profile.space_before)
    if profile.space_after is not None:
        paragraph_format.space_after = pt_or_none(profile.space_after)
    alignment = name_to_alignment(profile.alignment)
    if alignment is not None:
        paragraph_format.alignment = alignment


def _configure_toc_style(
    document: object,
    style_name: str,
    *,
    font_name: str,
    latin_font_name: str,
    font_size_pt: float,
    left_indent_pt: float,
    space_before_pt: float,
    space_after_pt: float,
) -> None:
    try:
        style = document.styles[style_name]
    except KeyError:
        return
    set_style_font(style, font_name, latin_font_name)
    style.font.size = Pt(font_size_pt)
    style.font.bold = False
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = Pt(20)
    paragraph_format.space_before = Pt(space_before_pt)
    paragraph_format.space_after = Pt(space_after_pt)
    paragraph_format.left_indent = Pt(left_indent_pt)
    paragraph_format.first_line_indent = Pt(0)


def _configure_north_college_toc_styles(document: object, template_profile: TemplateStyleProfile) -> None:
    if template_profile.profile_name != "hebei_north_college_fixed":
        return
    _configure_toc_style(
        document,
        "TOC 1",
        font_name="黑体",
        latin_font_name="Times New Roman",
        font_size_pt=12,
        left_indent_pt=0,
        space_before_pt=6,
        space_after_pt=0,
    )
    _configure_toc_style(
        document,
        "TOC 2",
        font_name="宋体",
        latin_font_name="Times New Roman",
        font_size_pt=12,
        left_indent_pt=24,
        space_before_pt=0,
        space_after_pt=0,
    )
    _configure_toc_style(
        document,
        "TOC 3",
        font_name="宋体",
        latin_font_name="Times New Roman",
        font_size_pt=12,
        left_indent_pt=48,
        space_before_pt=0,
        space_after_pt=0,
    )
    _configure_toc_style(
        document,
        "TOC 4",
        font_name="宋体",
        latin_font_name="Times New Roman",
        font_size_pt=12,
        left_indent_pt=72,
        space_before_pt=0,
        space_after_pt=0,
    )


def _format_north_college_toc_title(document: object, template_profile: TemplateStyleProfile) -> None:
    if template_profile.profile_name != "hebei_north_college_fixed":
        return
    for paragraph in document.paragraphs:
        if _compact_text(paragraph.text) == "目录":
            try:
                paragraph.style = document.styles["Heading 1"]
            except KeyError:
                pass
            _apply_paragraph_profile(paragraph, template_profile.heading_1)


def _apply_paragraph_profile(paragraph: object, profile: StyleProfile) -> None:
    paragraph_format = paragraph.paragraph_format
    if profile.line_spacing is not None:
        apply_line_spacing(paragraph_format, profile.line_spacing)
    if profile.first_line_indent is not None:
        paragraph_format.first_line_indent = pt_or_none(profile.first_line_indent)
    if profile.space_before is not None:
        paragraph_format.space_before = pt_or_none(profile.space_before)
    if profile.space_after is not None:
        paragraph_format.space_after = pt_or_none(profile.space_after)
    alignment = name_to_alignment(profile.alignment)
    if alignment is not None:
        paragraph.alignment = alignment

    for run in paragraph.runs:
        set_run_font(run, profile.font_name, profile.latin_font_name)
        font_size = font_size_to_pt(profile.font_size_pt)
        if font_size is not None:
            run.font.size = Pt(font_size)
        if profile.bold is not None:
            run.bold = profile.bold


def _replace_paragraph_text(paragraph: object, text: str) -> None:
    clear_paragraph_runs(paragraph)
    paragraph.add_run(text)


def _normalize_north_college_title_text(
    paragraph: object,
    prediction: ParagraphPrediction,
    template_profile: TemplateStyleProfile,
) -> None:
    if template_profile.profile_name != "hebei_north_college_fixed":
        return
    compact = _compact_text(paragraph.text)
    if prediction.predicted_type == "abstract_title":
        if normalize_text(paragraph.text).lower() == "abstract":
            _replace_paragraph_text(paragraph, "ABSTRACT")
        elif compact == "摘要":
            _replace_paragraph_text(paragraph, "摘  要")
    elif prediction.predicted_type == "contents_title" and compact == "目录":
        _replace_paragraph_text(paragraph, "目  录")
    elif prediction.predicted_type == "references_title" and compact != "参考文献":
        _replace_paragraph_text(paragraph, "参考文献")
    elif prediction.predicted_type == "keywords":
        text = normalize_text(paragraph.text)
        if re.match(r"^(?:keywords?|key\s+words?)\b", text, re.I):
            replaced = re.sub(r"^(?:keywords?|key\s+words?)", "Key Words", text, count=1, flags=re.I)
            _replace_paragraph_text(paragraph, replaced)
        elif compact.startswith("关键词") and not text.startswith("关键词"):
            suffix = compact[3:]
            _replace_paragraph_text(paragraph, f"关键词{suffix}")


def _apply_page_profile(document: object, profile: TemplateStyleProfile) -> None:
    for section in document.sections:
        top_margin = pt_or_none(profile.page.margin_top)
        bottom_margin = pt_or_none(profile.page.margin_bottom)
        left_margin = pt_or_none(profile.page.margin_left)
        right_margin = pt_or_none(profile.page.margin_right)
        if top_margin is not None:
            section.top_margin = top_margin
        if bottom_margin is not None:
            section.bottom_margin = bottom_margin
        if left_margin is not None:
            section.left_margin = left_margin
        if right_margin is not None:
            section.right_margin = right_margin


def _copy_header_footer(document: object, profile: TemplateStyleProfile) -> bool:
    if (
        not profile.header_footer.header_text
        and not profile.header_footer.footer_text
        and not profile.header_footer.has_page_number
    ):
        return False
    copied = True
    combined_text = f"{profile.header_footer.header_text or ''}\n{profile.header_footer.footer_text or ''}"
    already_has_visible_page_number = bool(re.search(r"PAGE|第\s*\d+\s*页|\d+\s*/\s*\d+", combined_text, re.I))
    for section in document.sections:
        ok = copy_header_footer_plain_text(
            section,
            profile.header_footer.header_text,
            profile.header_footer.footer_text,
        )
        if ok and profile.header_footer.has_page_number and not already_has_visible_page_number:
            footer_paragraphs = list(section.footer.paragraphs)
            if footer_paragraphs:
                footer_paragraphs[0].add_run(" ")
                add_page_number_field(footer_paragraphs[0])
        copied = copied and ok
    return copied


def _set_builtin_style(paragraph: object, predicted_type: str) -> None:
    style_name = BUILTIN_STYLE_BY_TYPE.get(predicted_type)
    if not style_name:
        return
    try:
        paragraph.style = style_name
    except Exception:
        pass


def _reference_body_indexes(predictions: list[ParagraphPrediction]) -> set[int]:
    reference_indexes: set[int] = set()
    in_references = False
    for prediction in predictions:
        if prediction.predicted_type == "references_title":
            in_references = True
            continue
        if not in_references:
            continue
        if prediction.predicted_type in {"appendix_title", "acknowledgements_title"}:
            break
        if prediction.predicted_type == "heading_1" and prediction.text.strip():
            break
        if prediction.predicted_type in {"body", "contents_entry"} and prediction.text.strip():
            reference_indexes.add(prediction.index)
    return reference_indexes


def _style_has_values(profile: StyleProfile) -> bool:
    return any(
        value is not None
        for key, value in profile.to_dict().items()
        if key not in {"outline_level", "numbering_pattern"}
    )


def _profile_for_type(
    template_profile: TemplateStyleProfile,
    effective_type: str,
    prediction_text: str,
    paragraph_index: int,
    reference_body_indexes: set[int],
) -> StyleProfile | None:
    if paragraph_index in reference_body_indexes:
        return template_profile.reference.body
    if effective_type == "reference_title":
        return template_profile.reference.title
    if effective_type == "abstract_title":
        if normalize_text(prediction_text).lower() == "abstract":
            return template_profile.english_abstract_title if _style_has_values(template_profile.english_abstract_title) else template_profile.heading_1
        return template_profile.abstract_title if _style_has_values(template_profile.abstract_title) else template_profile.heading_1
    if effective_type == "keywords":
        if re.match(r"^keywords?\b|^key\s+words?\b", normalize_text(prediction_text), re.I):
            return template_profile.english_keywords if _style_has_values(template_profile.english_keywords) else template_profile.body
        return template_profile.keywords if _style_has_values(template_profile.keywords) else template_profile.body
    profile_attr = PROFILE_BY_TYPE.get(effective_type)
    if profile_attr is None:
        return None
    if profile_attr == "reference_title":
        return template_profile.reference.title
    return getattr(template_profile, profile_attr)


def format_document(
    student_docx_path: str | Path,
    output_docx_path: str | Path,
    template_profile: TemplateStyleProfile,
    predictions: list[ParagraphPrediction],
    *,
    insert_toc: bool = True,
) -> FormatRunResult:
    student_path = Path(student_docx_path)
    document = Document(str(student_path))
    summary = FormatSummary()
    warnings: list[str] = []
    unsupported_objects = detect_complex_objects(student_path)

    _configure_document_style(document, "Normal", template_profile.body)
    _configure_document_style(document, "Heading 1", template_profile.heading_1)
    _configure_document_style(document, "Heading 2", template_profile.heading_2)
    _configure_document_style(document, "Heading 3", template_profile.heading_3)
    _configure_document_style(document, "Heading 4", template_profile.heading_4)
    _configure_north_college_toc_styles(document, template_profile)
    _apply_page_profile(document, template_profile)
    summary.header_footer_copied = _copy_header_footer(document, template_profile)
    has_header_footer_profile = bool(
        template_profile.header_footer.header_text
        or template_profile.header_footer.footer_text
        or template_profile.header_footer.has_page_number
    )
    if has_header_footer_profile and not summary.header_footer_copied:
        warnings.append("页眉页脚复制失败，已保留学生文档原始页眉页脚。")

    prediction_by_index = {prediction.index: prediction for prediction in predictions}
    reference_body_indexes = _reference_body_indexes(predictions)
    protected_until = _front_matter_end_index(predictions)
    north_front_matter_info = (
        extract_front_matter_info(document, protected_until)
        if template_profile.profile_name == "hebei_north_college_fixed"
        else None
    )
    if protected_until:
        warnings.append("已保护封面和声明区域：摘要前内容不套用正文/标题格式，避免破坏封面页。")
    summary.cover_heading_formatted = _format_cover_heading(document, protected_until)
    if summary.cover_heading_formatted:
        warnings.append("已单独修正封面主标题行距和段后距离。")
    summary.cover_title_formatted = _format_cover_title(document, protected_until, template_profile)
    if summary.cover_title_formatted:
        warnings.append("已单独修正封面论文题目行距，避免长标题换行后重叠。")
    removed_cover_blanks = _compact_cover_spacing(document, protected_until)
    if removed_cover_blanks:
        warnings.append(f"已压缩封面主标题和论文题目之间的多余空段 {removed_cover_blanks} 个，封面应保持单页。")
    cover_metadata_count = _format_cover_metadata(document, protected_until, template_profile)
    if cover_metadata_count:
        warnings.append(f"已按河北北方学院封面要求检查并格式化 {cover_metadata_count} 个封面信息段落。")
    declaration_count = _format_declaration_signature_area(document, protected_until, template_profile)
    if declaration_count:
        warnings.append("已按原创性声明页要求修正签名区和日期行的宋体小四格式及位置。")
    uncertain = [
        prediction
        for prediction in predictions
        if prediction.predicted_type.startswith("heading") and prediction.confidence < 0.75
    ]
    summary.uncertain_paragraphs = len(uncertain)

    for index, paragraph in enumerate(document.paragraphs):
        prediction = prediction_by_index.get(index)
        if prediction is None:
            continue
        effective_type = _effective_prediction_type(prediction, protected_until)
        if effective_type is None:
            continue
        profile = _profile_for_type(template_profile, effective_type, prediction.text, index, reference_body_indexes)
        if profile is None:
            continue
        _normalize_north_college_title_text(paragraph, prediction, template_profile)
        _set_builtin_style(paragraph, effective_type)
        _apply_paragraph_profile(paragraph, profile)
        summary_field = SUMMARY_FIELD_BY_TYPE.get(effective_type)
        if summary_field:
            setattr(summary, summary_field, getattr(summary, summary_field) + 1)
        elif effective_type != "keywords":
            summary.special_paragraphs_modified += 1

    table_result = format_all_tables(document, template_profile)
    summary.tables_detected = table_result.tables_detected
    summary.tables_formatted = table_result.tables_formatted
    summary.three_line_tables_applied = table_result.three_line_tables_applied
    summary.table_captions_detected = table_result.table_captions_detected
    summary.table_captions_formatted = table_result.table_captions_formatted
    unsupported_objects.extend(table_result.complex_tables)
    warnings.extend(table_result.warnings)

    figure_result = format_all_figures(document, template_profile)
    summary.figures_detected = figure_result.figures_detected
    summary.figures_centered = figure_result.figures_centered
    summary.figure_captions_detected = figure_result.figure_captions_detected
    summary.figure_captions_formatted = figure_result.figure_captions_formatted
    unsupported_objects.extend(figure_result.uncertain_links)
    warnings.extend(figure_result.warnings)

    toc_title = "目  录" if template_profile.profile_name == "hebei_north_college_fixed" else "目录"
    rebuild_toc = template_profile.profile_name == "hebei_north_college_fixed"
    toc_result = insert_or_mark_toc(document, predictions, title_text=toc_title, rebuild_existing=rebuild_toc) if insert_toc else None
    _format_north_college_toc_title(document, template_profile)
    if toc_result is None:
        from app.models import TocResult

        toc_result = TocResult(False, False, "未请求插入目录。")
    summary.toc_inserted = toc_result.inserted or toc_result.existing_toc_detected
    if not toc_result.inserted and not toc_result.existing_toc_detected:
        warnings.append(toc_result.message)

    if north_front_matter_info and replace_north_college_front_matter(document, protected_until, north_front_matter_info):
        warnings.append("已使用北方学院固定封面/声明页模板替换学生论文原前置页，仅填入论文题目和可识别的学生信息。")
        summary.cover_heading_formatted = True
        summary.cover_title_formatted = True

    output_path = Path(output_docx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return FormatRunResult(
        summary=summary,
        toc=toc_result,
        uncertain_predictions=uncertain,
        unsupported_objects=unsupported_objects,
        warnings=warnings,
    )
