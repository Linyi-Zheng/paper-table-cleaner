from __future__ import annotations

from app.models import ParagraphPrediction, TocResult
from app.services.ooxml_utils import (
    add_toc_field,
    has_toc_field,
    insert_paragraph_after,
    insert_paragraph_before,
    normalize_text,
    set_update_fields_on_open,
)


def _remove_paragraph(paragraph: object) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
        paragraph._p = paragraph._element = None


def _looks_like_manual_toc_entry(paragraph: object) -> bool:
    style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
    text = normalize_text(paragraph.text)
    if style_name.startswith("toc"):
        return True
    if _paragraph_has_toc_field(paragraph):
        return True
    if "\t" in paragraph.text and text:
        return True
    return False


def _paragraph_has_toc_field(paragraph: object) -> bool:
    for node in paragraph._element.iter():
        if node.tag.endswith("}instrText") and node.text and "TOC" in node.text.upper():
            return True
    return False


def _remove_all_toc_field_paragraphs(document: object) -> int:
    removed = 0
    for paragraph in list(document.paragraphs):
        if _paragraph_has_toc_field(paragraph):
            _remove_paragraph(paragraph)
            removed += 1
    return removed


def _remove_manual_toc_entries_after(contents_title: object) -> int:
    removed = 0
    current = contents_title._element.getnext()
    while current is not None:
        from docx.text.paragraph import Paragraph

        paragraph = Paragraph(current, contents_title._parent)
        text = normalize_text(paragraph.text)
        if not text:
            next_element = current.getnext()
            _remove_paragraph(paragraph)
            removed += 1
            current = next_element
            continue
        if not _looks_like_manual_toc_entry(paragraph):
            break
        next_element = current.getnext()
        _remove_paragraph(paragraph)
        removed += 1
        current = next_element
    return removed


def _find_contents_title(document: object):
    for paragraph in document.paragraphs:
        if normalize_text(paragraph.text).replace(" ", "") == "目录":
            return paragraph
    return None


def _find_abstract_end(document: object):
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text).lower()
        if text in {"关键词", "关键字"} or text.startswith(("关键词", "关键字", "keywords", "key words")):
            return paragraph
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if text.replace(" ", "") == "摘要" or text.lower() == "abstract":
            return paragraph
    return None


def _find_first_heading(document: object, predictions: list[ParagraphPrediction]):
    heading_indexes = {
        prediction.index
        for prediction in predictions
        if prediction.predicted_type in {"heading_1", "heading_2", "heading_3"}
    }
    for index, paragraph in enumerate(document.paragraphs):
        if index in heading_indexes:
            return paragraph
    return document.paragraphs[0] if document.paragraphs else None


def insert_or_mark_toc(
    document: object,
    predictions: list[ParagraphPrediction],
    *,
    title_text: str = "目录",
    rebuild_existing: bool = False,
) -> TocResult:
    existing_toc = has_toc_field(document)
    if existing_toc and not rebuild_existing:
        set_update_fields_on_open(document)
        return TocResult(
            inserted=False,
            existing_toc_detected=True,
            message="检测到已有目录字段，已设置为打开文档时更新域。",
        )

    try:
        removed_existing_fields = _remove_all_toc_field_paragraphs(document) if rebuild_existing and existing_toc else 0
        contents_title = _find_contents_title(document)
        if contents_title is not None:
            removed_count = _remove_manual_toc_entries_after(contents_title)
            toc_paragraph = insert_paragraph_after(contents_title)
            add_toc_field(toc_paragraph)
            message = "已插入 Word TOC 字段，打开 Word 后需要更新域以刷新页码。"
            if removed_existing_fields:
                message = f"已重建旧目录字段 {removed_existing_fields} 处，并插入新的 Word TOC 字段；打开 Word 后需要更新域。"
            if removed_count:
                message = f"已移除 {removed_count} 条旧手工目录，并插入 Word TOC 字段；打开 Word 后需要更新域。"
        else:
            anchor = _find_abstract_end(document) or _find_first_heading(document, predictions)
            if anchor is None:
                return TocResult(False, False, "文档为空，无法插入目录。")
            if _find_abstract_end(document) is not None:
                title = insert_paragraph_after(anchor, title_text, "Heading 1")
                toc_paragraph = insert_paragraph_after(title)
            else:
                title = insert_paragraph_before(anchor, title_text, "Heading 1")
                toc_paragraph = insert_paragraph_after(title)
            add_toc_field(toc_paragraph)
            message = "已插入 Word TOC 字段，打开 Word 后需要更新域以刷新页码。"
        set_update_fields_on_open(document)
        return TocResult(inserted=True, existing_toc_detected=False, message=message)
    except Exception as exc:
        return TocResult(
            inserted=False,
            existing_toc_detected=False,
            message=f"目录插入失败：{exc}",
        )
