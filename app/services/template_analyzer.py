from __future__ import annotations

from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from app.models import HeaderFooterProfile, PageProfile, StyleProfile, TemplateStyleProfile
from app.services.ooxml_utils import (
    alignment_to_name,
    detect_page_number,
    font_size_to_pt,
    length_to_pt,
    normalize_line_spacing,
    normalize_text,
)
from app.services.figure_formatter import detect_figure_caption_style_from_template
from app.services.student_analyzer import extract_features, classify_text
from app.services.table_formatter import detect_table_style_from_template


def _font_name_from_style(style: object) -> str | None:
    font_name = getattr(getattr(style, "font", None), "name", None)
    if font_name:
        return font_name
    base_style = getattr(style, "base_style", None)
    if base_style is not None and base_style is not style:
        return _font_name_from_style(base_style)
    return None


def _font_size_from_style(style: object) -> float | None:
    size = getattr(getattr(style, "font", None), "size", None)
    if size is not None:
        return font_size_to_pt(size)
    base_style = getattr(style, "base_style", None)
    if base_style is not None and base_style is not style:
        return _font_size_from_style(base_style)
    return None


def _bold_from_style(style: object) -> bool | None:
    bold = getattr(getattr(style, "font", None), "bold", None)
    if bold is not None:
        return bool(bold)
    base_style = getattr(style, "base_style", None)
    if base_style is not None and base_style is not style:
        return _bold_from_style(base_style)
    return None


def _paragraph_profile_from_style(style: object, outline_level: int | None = None) -> StyleProfile:
    paragraph_format = style.paragraph_format
    return StyleProfile(
        font_name=_font_name_from_style(style),
        font_size_pt=_font_size_from_style(style),
        bold=_bold_from_style(style),
        line_spacing=_line_spacing_to_value(paragraph_format.line_spacing),
        first_line_indent=length_to_pt(paragraph_format.first_line_indent),
        alignment=alignment_to_name(paragraph_format.alignment),
        space_before=length_to_pt(paragraph_format.space_before),
        space_after=length_to_pt(paragraph_format.space_after),
        outline_level=outline_level,
    )


def _line_spacing_to_value(value: object | None) -> float | None:
    return normalize_line_spacing(value)


def _profile_from_paragraph(paragraph: object, outline_level: int | None = None) -> StyleProfile:
    features = extract_features(paragraph)
    paragraph_format = paragraph.paragraph_format
    font_names = [
        run.font.name
        for run in paragraph.runs
        if run.text.strip() and run.font.name
    ]
    sizes = [
        font_size_to_pt(run.font.size)
        for run in paragraph.runs
        if run.text.strip() and run.font.size is not None
    ]
    sizes = [size for size in sizes if size is not None]
    style_profile = _paragraph_profile_from_style(paragraph.style, outline_level)
    return StyleProfile(
        font_name=Counter(font_names).most_common(1)[0][0] if font_names else style_profile.font_name,
        font_size_pt=Counter(sizes).most_common(1)[0][0] if sizes else style_profile.font_size_pt,
        bold=features.has_bold_run if features.has_bold_run else style_profile.bold,
        line_spacing=_line_spacing_to_value(paragraph_format.line_spacing) or style_profile.line_spacing,
        first_line_indent=length_to_pt(paragraph_format.first_line_indent) or style_profile.first_line_indent,
        alignment=alignment_to_name(paragraph.alignment) or style_profile.alignment,
        space_before=length_to_pt(paragraph_format.space_before) or style_profile.space_before,
        space_after=length_to_pt(paragraph_format.space_after) or style_profile.space_after,
        outline_level=outline_level,
    )


def _get_style(document: object, names: list[str]) -> object | None:
    for name in names:
        try:
            style = document.styles[name]
        except KeyError:
            continue
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            return style
    return None


def _infer_profile_from_paragraphs(
    document: object,
    predicted_type: str,
    outline_level: int | None,
) -> StyleProfile | None:
    candidates: list[tuple[float, object]] = []
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        features = extract_features(paragraph)
        prediction = classify_text(
            text,
            style_name=features.style_name,
            is_centered=features.is_centered,
            has_bold_run=features.has_bold_run,
            max_font_size_pt=features.max_font_size_pt,
        )
        if prediction.predicted_type == predicted_type:
            candidates.append((prediction.confidence, paragraph))
    if not candidates:
        return None
    candidates.sort(key=lambda item: item[0], reverse=True)
    return _profile_from_paragraph(candidates[0][1], outline_level)


def _infer_body_profile(document: object) -> StyleProfile | None:
    candidates = [
        paragraph
        for paragraph in document.paragraphs
        if len(normalize_text(paragraph.text)) >= 30
        and classify_text(normalize_text(paragraph.text)).predicted_type == "body"
    ]
    if not candidates:
        return None
    return _profile_from_paragraph(candidates[0], None)


def _page_profile(document: object) -> PageProfile:
    section = document.sections[0]
    return PageProfile(
        margin_top=length_to_pt(section.top_margin),
        margin_bottom=length_to_pt(section.bottom_margin),
        margin_left=length_to_pt(section.left_margin),
        margin_right=length_to_pt(section.right_margin),
    )


def _header_footer_profile(document: object) -> HeaderFooterProfile:
    section = document.sections[0]
    header_text = "\n".join(
        normalize_text(paragraph.text)
        for paragraph in section.header.paragraphs
        if normalize_text(paragraph.text)
    )
    footer_text = "\n".join(
        normalize_text(paragraph.text)
        for paragraph in section.footer.paragraphs
        if normalize_text(paragraph.text)
    )
    return HeaderFooterProfile(
        header_text=header_text or None,
        footer_text=footer_text or None,
        has_page_number=detect_page_number(section.footer) or detect_page_number(section.header),
    )


def _missing_fields(profile: TemplateStyleProfile) -> list[str]:
    missing: list[str] = []
    sections = {
        "body": profile.body,
        "heading_1": profile.heading_1,
        "heading_2": profile.heading_2,
        "heading_3": profile.heading_3,
    }
    for name, style in sections.items():
        for field_name, value in style.to_dict().items():
            if field_name in {"outline_level", "numbering_pattern"}:
                continue
            if value is None:
                missing.append(f"{name}.{field_name}")
    for field_name, value in profile.page.to_dict().items():
        if value is None:
            missing.append(f"page.{field_name}")
    if profile.header_footer.has_page_number is None:
        missing.append("header_footer.has_page_number")
    return missing


def analyze_template(docx_path: str | Path) -> TemplateStyleProfile:
    document = Document(str(docx_path))
    profile = TemplateStyleProfile()

    style_map = {
        "body": (["Normal", "正文"], None),
        "heading_1": (["Heading 1", "标题 1", "标题1"], 1),
        "heading_2": (["Heading 2", "标题 2", "标题2"], 2),
        "heading_3": (["Heading 3", "标题 3", "标题3"], 3),
    }

    for key, (names, outline_level) in style_map.items():
        style = _get_style(document, names)
        extracted = _paragraph_profile_from_style(style, outline_level) if style is not None else None
        if extracted is None or extracted.font_size_pt is None:
            inferred = _infer_body_profile(document) if key == "body" else _infer_profile_from_paragraphs(document, key, outline_level)
            if inferred is not None:
                extracted = inferred
                profile.warnings.append(f"{key} 未能完全从内置样式读取，已使用模板段落启发式推断。")
        if extracted is not None:
            setattr(profile, key, extracted)
        else:
            profile.warnings.append(f"未找到 {key} 的可靠样式。")

    profile.heading_1.outline_level = 1
    profile.heading_2.outline_level = 2
    profile.heading_3.outline_level = 3
    profile.page = _page_profile(document)
    profile.header_footer = _header_footer_profile(document)
    profile.table = detect_table_style_from_template(document)
    profile.figure = detect_figure_caption_style_from_template(document)
    if profile.table.used_default:
        profile.warnings.append("模板中未识别到明确表格样式，后续将使用默认三线表样式。")
    if profile.figure.used_default:
        profile.warnings.append("模板中未识别到明确图题样式，后续将使用默认图题样式。")
    profile.missing_fields = _missing_fields(profile)
    if profile.missing_fields:
        profile.warnings.append("部分模板样式字段无法确定，报告中已列出 null 字段。")
    return profile
