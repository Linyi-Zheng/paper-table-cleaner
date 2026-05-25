from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.models import FormatRunResult, ParagraphPrediction, TemplateStyleProfile
from app.services.ooxml_utils import set_run_font, set_style_font


def _markdown_style_block(profile: TemplateStyleProfile) -> str:
    data = profile.to_dict()
    return "```json\n" + json.dumps(data, ensure_ascii=False, indent=2) + "\n```"


def _prediction_table(predictions: list[ParagraphPrediction]) -> str:
    if not predictions:
        return "无。"
    lines = ["| 段落序号 | 类型 | 置信度 | 文本 |", "| --- | --- | --- | --- |"]
    for item in predictions[:50]:
        text = item.text.replace("|", "\\|")
        if len(text) > 80:
            text = text[:77] + "..."
        lines.append(f"| {item.index} | {item.predicted_type} | {item.confidence:.2f} | {text} |")
    if len(predictions) > 50:
        lines.append(f"| ... | ... | ... | 另有 {len(predictions) - 50} 条未展示 |")
    return "\n".join(lines)


def build_report_data(
    template_profile: TemplateStyleProfile,
    predictions: list[ParagraphPrediction],
    format_result: FormatRunResult,
) -> dict[str, object]:
    heading_count = sum(1 for item in predictions if item.predicted_type.startswith("heading"))
    return {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "template_profile": template_profile.to_dict(),
        "summary": format_result.summary.to_dict(),
        "student_paragraphs_total": len(predictions),
        "student_headings_detected": heading_count,
        "uncertain_paragraphs": [item.to_dict() for item in format_result.uncertain_predictions],
        "unsupported_objects": format_result.unsupported_objects,
        "toc": format_result.toc.to_dict(),
        "warnings": template_profile.warnings + format_result.warnings,
        "table_template": {
            "is_three_line": template_profile.table.is_three_line,
            "used_default": template_profile.table.used_default,
        },
        "figure_template": {
            "image_alignment": template_profile.figure.image_alignment,
            "used_default": template_profile.figure.used_default,
        },
        "manual_actions_required": [
            "请使用 Microsoft Word 或 WPS 打开输出文档并更新目录域，以刷新目录页码。",
            "提交前请人工复核低置信度标题层级、表题、图题和学校特殊格式要求。",
        ],
    }


def build_markdown_report(
    template_profile: TemplateStyleProfile,
    predictions: list[ParagraphPrediction],
    format_result: FormatRunResult,
) -> str:
    summary = format_result.summary
    warnings = template_profile.warnings + format_result.warnings
    unsupported = format_result.unsupported_objects
    lines = [
        "# 论文格式修正报告",
        "",
        "## 处理摘要",
        "",
        f"- 正文段落修改：{summary.body_paragraphs_modified}",
        f"- 一级标题修改：{summary.heading_1_modified}",
        f"- 二级标题修改：{summary.heading_2_modified}",
        f"- 三级标题修改：{summary.heading_3_modified}",
        f"- 四级标题修改：{summary.heading_4_modified}",
        f"- 特殊段落修改：{summary.special_paragraphs_modified}",
        f"- 表格检测：{summary.tables_detected}",
        f"- 表格格式化：{summary.tables_formatted}",
        f"- 三线表应用：{summary.three_line_tables_applied}",
        f"- 表题检测：{summary.table_captions_detected}",
        f"- 表题格式化：{summary.table_captions_formatted}",
        f"- 图片检测：{summary.figures_detected}",
        f"- 图片居中：{summary.figures_centered}",
        f"- 图题检测：{summary.figure_captions_detected}",
        f"- 图题格式化：{summary.figure_captions_formatted}",
        f"- 封面主标题行距修正：{'是' if summary.cover_heading_formatted else '否'}",
        f"- 封面题目行距修正：{'是' if summary.cover_title_formatted else '否'}",
        f"- 低置信度标题：{summary.uncertain_paragraphs}",
        f"- 目录字段已插入：{'是' if summary.toc_inserted else '否'}",
        f"- 页眉页脚已复制：{'是' if summary.header_footer_copied else '否'}",
        "",
        "## 图表样式判断",
        "",
        f"- 模板识别到三线表：{'是' if template_profile.table.is_three_line else '否'}",
        f"- 表格样式来源：{'默认中文论文样式' if template_profile.table.used_default else '模板提取'}",
        f"- 图题样式来源：{'默认中文论文样式' if template_profile.figure.used_default else '模板提取'}",
        "",
        "## 模板样式识别结果",
        "",
        _markdown_style_block(template_profile),
        "",
        "## 低置信度标题层级",
        "",
        _prediction_table(format_result.uncertain_predictions),
        "",
        "## 无法自动处理或仅保留的内容",
        "",
    ]
    if unsupported:
        lines.extend(f"- {item}" for item in unsupported)
    else:
        lines.append("- 未检测到需要特别提示的复杂对象。")
    lines.extend(
        [
            "",
            "## 目录处理",
            "",
            f"- {format_result.toc.message}",
            "- 目录页码可能需要打开 Word 后右键更新域。",
            "",
            "## 提示与限制",
            "",
        ]
    )
    if warnings:
        lines.extend(f"- {item}" for item in warnings)
    else:
        lines.append("- 未发现额外警告。")
    lines.extend(
        [
            "- 当前 MVP 仅支持 .docx。",
            "- 对复杂页眉页脚、公式、嵌入对象、浮动图片和复杂合并单元格支持有限。",
            "- 表格编号、图表编号和题注文字内容不会被强制改写。",
            "- 模板样式不规范时会使用启发式推断和默认中文论文图表格式。",
            "- 输出文档请由学生最终检查后提交。",
            "",
        ]
    )
    return "\n".join(lines)


def _add_report_paragraph(document: object, text: str, *, bold_label: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", "Times New Roman")
    run.font.size = Pt(10.5)
    if bold_label:
        run.bold = True


def _add_report_heading(document: object, text: str, level: int = 1) -> None:
    paragraph = document.add_heading("", level=level)
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", "Times New Roman")
    run.bold = True
    run.font.size = Pt(15 if level == 1 else 12)
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_word_report(
    template_profile: TemplateStyleProfile,
    predictions: list[ParagraphPrediction],
    format_result: FormatRunResult,
) -> Document:
    document = Document()
    set_style_font(document.styles["Normal"], "宋体", "Times New Roman")
    document.styles["Normal"].font.size = Pt(10.5)

    data = build_report_data(template_profile, predictions, format_result)
    summary = format_result.summary
    warnings = template_profile.warnings + format_result.warnings

    _add_report_heading(document, "论文格式修正报告", 1)
    _add_report_paragraph(document, f"生成时间：{data['generated_at']}")
    _add_report_paragraph(document, f"格式模块：{template_profile.profile_name}")

    _add_report_heading(document, "处理摘要", 2)
    summary_rows = [
        ("正文段落修改", summary.body_paragraphs_modified),
        ("一级标题修改", summary.heading_1_modified),
        ("二级标题修改", summary.heading_2_modified),
        ("三级标题修改", summary.heading_3_modified),
        ("四级标题修改", summary.heading_4_modified),
        ("特殊段落修改", summary.special_paragraphs_modified),
        ("表格检测/格式化", f"{summary.tables_detected}/{summary.tables_formatted}"),
        ("三线表应用", summary.three_line_tables_applied),
        ("表题检测/格式化", f"{summary.table_captions_detected}/{summary.table_captions_formatted}"),
        ("图片检测/居中", f"{summary.figures_detected}/{summary.figures_centered}"),
        ("图题检测/格式化", f"{summary.figure_captions_detected}/{summary.figure_captions_formatted}"),
        ("封面主标题行距修正", "是" if summary.cover_heading_formatted else "否"),
        ("封面题目行距修正", "是" if summary.cover_title_formatted else "否"),
        ("低置信度标题", summary.uncertain_paragraphs),
        ("目录已处理", "是" if summary.toc_inserted else "否"),
        ("页眉页脚已处理", "是" if summary.header_footer_copied else "否"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "结果"
    for label, value in summary_rows:
        row = table.add_row()
        row.cells[0].text = str(label)
        row.cells[1].text = str(value)

    _add_report_heading(document, "图表样式判断", 2)
    _add_report_paragraph(document, f"模板识别到三线表：{'是' if template_profile.table.is_three_line else '否'}")
    _add_report_paragraph(document, f"表格样式来源：{'默认中文论文样式' if template_profile.table.used_default else '固定规则或模板提取'}")
    _add_report_paragraph(document, f"图题样式来源：{'默认中文论文样式' if template_profile.figure.used_default else '固定规则或模板提取'}")

    _add_report_heading(document, "低置信度标题层级", 2)
    if format_result.uncertain_predictions:
        for item in format_result.uncertain_predictions[:50]:
            text = item.text[:100] + ("..." if len(item.text) > 100 else "")
            _add_report_paragraph(document, f"段落 {item.index}：{item.predicted_type}，置信度 {item.confidence:.2f}，{text}")
    else:
        _add_report_paragraph(document, "无。")

    _add_report_heading(document, "无法自动处理或仅保留的内容", 2)
    if format_result.unsupported_objects:
        for item in format_result.unsupported_objects:
            _add_report_paragraph(document, item)
    else:
        _add_report_paragraph(document, "未检测到需要特别提示的复杂对象。")

    _add_report_heading(document, "目录处理", 2)
    _add_report_paragraph(document, format_result.toc.message)
    _add_report_paragraph(document, "目录页码可能需要打开 Word 后右键更新域。")

    _add_report_heading(document, "提示与限制", 2)
    for item in warnings:
        _add_report_paragraph(document, item)
    _add_report_paragraph(document, "当前 MVP 仅支持 .docx。")
    _add_report_paragraph(document, "对复杂页眉页脚、公式、嵌入对象、浮动图片和复杂合并单元格支持有限。")
    _add_report_paragraph(document, "输出文档请由学生最终检查后提交。")
    return document


def write_reports(
    template_profile: TemplateStyleProfile,
    predictions: list[ParagraphPrediction],
    format_result: FormatRunResult,
    report_json_path: str | Path,
    report_path: str | Path,
    report_md_path: str | Path | None = None,
) -> dict[str, object]:
    data = build_report_data(template_profile, predictions, format_result)
    json_path = Path(report_json_path)
    output_path = Path(report_path)
    json_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    if output_path.suffix.lower() == ".docx":
        output_path.parent.mkdir(parents=True, exist_ok=True)
        build_word_report(template_profile, predictions, format_result).save(str(output_path))
    else:
        output_path.write_text(
            build_markdown_report(template_profile, predictions, format_result),
            encoding="utf-8",
        )
    if report_md_path is not None:
        Path(report_md_path).write_text(
            build_markdown_report(template_profile, predictions, format_result),
            encoding="utf-8",
        )
    return data
