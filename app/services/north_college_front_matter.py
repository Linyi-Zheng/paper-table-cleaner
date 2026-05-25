from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn

from app.services.ooxml_utils import normalize_text

ASSET_PATH = Path(__file__).resolve().parents[1] / "assets" / "north_college_front_matter.docx"
FRONT_MATTER_LAST_PARAGRAPH_TEXT = "年   月   日                      年   月   日"


@dataclass(slots=True)
class FrontMatterInfo:
    title: str | None = None
    author: str | None = None
    student_id: str | None = None
    department: str | None = None
    major: str | None = None
    advisor: str | None = None
    date_text: str | None = None


def _compact(text: str) -> str:
    return re.sub(r"\s+", "", normalize_text(text))


def _paragraph_text(paragraph: object) -> str:
    return normalize_text(getattr(paragraph, "text", ""))


def _looks_like_cover_title(paragraph: object) -> bool:
    text = _paragraph_text(paragraph)
    compact = _compact(text)
    if not compact or len(compact) < 10:
        return False
    if any(label in compact for label in ("学士学位论文", "姓名", "学号", "院系", "专业", "指导教师", "年月日")):
        return False
    if any(word in compact for word in ("声明", "授权书", "论文作者", "指导教师确认", "本人所提交")):
        return False
    if re.match(r"^\d+(?:\.\d+)*", compact):
        return False
    return True


def _value_after_label(text: str, label: str) -> str | None:
    compact = _compact(text)
    if not compact.startswith(label):
        return None
    value = compact[len(label) :].strip("：: _-—")
    return value or None


def extract_front_matter_info(document: object, protected_until: int) -> FrontMatterInfo:
    scan_until = protected_until if protected_until > 0 else min(80, len(document.paragraphs))
    paragraphs = list(document.paragraphs[:scan_until])
    info = FrontMatterInfo()

    title_candidates = [paragraph for paragraph in paragraphs if _looks_like_cover_title(paragraph)]
    if title_candidates:
        info.title = _paragraph_text(max(title_candidates, key=lambda paragraph: len(_compact(paragraph.text))))

    for paragraph in paragraphs:
        text = _paragraph_text(paragraph)
        if not text:
            continue
        info.author = info.author or _value_after_label(text, "姓名")
        info.student_id = info.student_id or _value_after_label(text, "学号")
        info.department = info.department or _value_after_label(text, "院系")
        info.major = info.major or _value_after_label(text, "专业")
        info.advisor = info.advisor or _value_after_label(text, "指导教师")
        if info.date_text is None and re.search(r"[一二三四五六七八九十〇零\d]{4}年.*月.*日", text):
            info.date_text = text
    return info


def _set_paragraph_text_preserve_first_run(paragraph: object, text: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def _first_underlined_run(paragraph: object):
    for run in paragraph.runs:
        if run.underline:
            return run
    return None


def _set_underlined_value(paragraph: object, value: str | None) -> None:
    if not value:
        return
    run = _first_underlined_run(paragraph)
    if run is None:
        paragraph.add_run(value)
        return
    original_width = max(len(run.text), len(value) + 4)
    run.text = value.center(original_width)


def _set_cover_title(paragraph: object, title: str) -> None:
    run = _first_underlined_run(paragraph)
    if run is None:
        _set_paragraph_text_preserve_first_run(paragraph, title)
    else:
        run.text = title
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None


def _fill_template(template_document: object, info: FrontMatterInfo) -> None:
    title = info.title or ""
    for paragraph in template_document.paragraphs:
        text = paragraph.text
        compact = _compact(text)
        if title and text and not compact and len(text) >= 8:
            _set_cover_title(paragraph, title)
        elif compact.startswith("姓名"):
            _set_underlined_value(paragraph, info.author)
        elif compact.startswith("学号"):
            _set_underlined_value(paragraph, info.student_id)
        elif compact.startswith("院系"):
            _set_underlined_value(paragraph, info.department)
        elif compact.startswith("专业"):
            _set_underlined_value(paragraph, info.major)
        elif compact.startswith("指导教师") and "确认" not in compact:
            _set_underlined_value(paragraph, info.advisor)
        elif compact == "年月日" and info.date_text:
            _set_paragraph_text_preserve_first_run(paragraph, info.date_text)
        elif title and text.startswith("本人所提交的学位论文"):
            _set_paragraph_text_preserve_first_run(
                paragraph,
                f"本人所提交的学位论文《{title}》，是在导师的指导下，独立进行研究工作所取得的原创性成果。"
                "除文中已经注明引用的内容外，本论文不包含任何其他个人或集体已经发表或撰写过的研究成果。"
                "对本文的研究做出重要贡献的个人和集体，均已在文中标明。",
            )


def _front_matter_elements(template_document: object) -> list[object]:
    elements: list[object] = []
    for child in template_document.element.body:
        if child.tag == qn("w:sectPr"):
            continue
        elements.append(deepcopy(child))
        text = "".join(t.text or "" for t in child.iter(qn("w:t")))
        if normalize_text(text) == FRONT_MATTER_LAST_PARAGRAPH_TEXT:
            break
    return elements


def _image_relationship_map(template_document: object, target_document: object) -> dict[str, str]:
    rel_map: dict[str, str] = {}
    for rel_id, rel in template_document.part.rels.items():
        if rel.reltype != RT.IMAGE:
            continue
        image_stream = BytesIO(rel.target_part.blob)
        new_rel_id, _ = target_document.part.get_or_add_image(image_stream)
        rel_map[rel_id] = new_rel_id
    return rel_map


def _remap_relationships(elements: list[object], rel_map: dict[str, str]) -> None:
    if not rel_map:
        return
    for element in elements:
        for node in element.iter():
            for attr_name, attr_value in list(node.attrib.items()):
                if attr_value in rel_map and (
                    attr_name == qn("r:embed") or attr_name == qn("r:link") or attr_name.endswith("}id")
                ):
                    node.set(attr_name, rel_map[attr_value])


def _remove_existing_front_matter(document: object, protected_until: int) -> object | None:
    if protected_until <= 0 or protected_until >= len(document.paragraphs):
        return None
    anchor = document.paragraphs[protected_until]._element
    body = document.element.body
    for child in list(body):
        if child is anchor:
            break
        body.remove(child)
    return anchor


def replace_north_college_front_matter(document: object, protected_until: int, info: FrontMatterInfo) -> bool:
    if not ASSET_PATH.exists():
        return False
    anchor = _remove_existing_front_matter(document, protected_until)
    if anchor is None:
        return False

    template_document = Document(str(ASSET_PATH))
    _fill_template(template_document, info)
    elements = _front_matter_elements(template_document)
    rel_map = _image_relationship_map(template_document, document)
    _remap_relationships(elements, rel_map)

    for element in elements:
        anchor.addprevious(element)
    return True
