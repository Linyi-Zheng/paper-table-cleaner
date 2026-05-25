from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass, field

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.table import Table

from app.models import StyleProfile, TableProfile
from app.services.ooxml_utils import (
    alignment_to_name,
    apply_line_spacing,
    clear_paragraph_runs,
    clear_table_style_and_fills,
    apply_three_line_borders,
    border_value,
    font_size_to_pt,
    length_to_pt,
    name_to_alignment,
    normalize_line_spacing,
    normalize_text,
    pt_or_none,
    set_run_font,
    table_has_merged_cells,
)

TABLE_CAPTION_RE = re.compile(
    r"^(表\s*\d+(?:[-.－—]\d+)?\s*\S+|表\s*[一二三四五六七八九十]+\s*\S+|Table\s+\d+(?:[-.]\d+)?\s*\S+)",
    re.I,
)
TABLE_NOTE_RE = re.compile(r"^(注[:：]|表注[:：]|资料来源[:：]|数据来源[:：])")
TABLE_CAPTION_PREFIX_RE = re.compile(
    r"^(表)\s*([0-9一二三四五六七八九十]+(?:[-.－—][0-9一二三四五六七八九十]+)?)\s*(.*)$",
    re.I,
)
ENGLISH_TABLE_CAPTION_PREFIX_RE = re.compile(r"^(Table)\s+([0-9]+(?:[-.][0-9]+)?)\s*(.*)$", re.I)


@dataclass(slots=True)
class TableFormatResult:
    tables_detected: int = 0
    tables_formatted: int = 0
    three_line_tables_applied: int = 0
    table_captions_detected: int = 0
    table_captions_formatted: int = 0
    complex_tables: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def is_table_caption_text(text: str) -> bool:
    return bool(TABLE_CAPTION_RE.match(normalize_text(text)))


def find_table_captions(document: object) -> list[object]:
    return [paragraph for paragraph in document.paragraphs if is_table_caption_text(paragraph.text)]


def is_table_note_text(text: str) -> bool:
    return bool(TABLE_NOTE_RE.match(normalize_text(text)))


def find_table_notes(document: object) -> list[object]:
    return [paragraph for paragraph in document.paragraphs if is_table_note_text(paragraph.text)]


def _profile_from_paragraph(paragraph: object, fallback: StyleProfile | None = None) -> StyleProfile:
    fallback = fallback or StyleProfile()
    font_names = [run.font.name for run in paragraph.runs if run.text.strip() and run.font.name]
    sizes = [font_size_to_pt(run.font.size) for run in paragraph.runs if run.text.strip() and run.font.size]
    sizes = [size for size in sizes if size is not None]
    bold_values = [run.bold for run in paragraph.runs if run.text.strip() and run.bold is not None]
    paragraph_format = paragraph.paragraph_format
    return StyleProfile(
        font_name=Counter(font_names).most_common(1)[0][0] if font_names else fallback.font_name,
        font_size_pt=Counter(sizes).most_common(1)[0][0] if sizes else fallback.font_size_pt,
        bold=Counter(bold_values).most_common(1)[0][0] if bold_values else fallback.bold,
        line_spacing=normalize_line_spacing(paragraph_format.line_spacing) or fallback.line_spacing,
        first_line_indent=length_to_pt(paragraph_format.first_line_indent) or fallback.first_line_indent,
        alignment=alignment_to_name(paragraph.alignment) or fallback.alignment,
        space_before=length_to_pt(paragraph_format.space_before) or fallback.space_before,
        space_after=length_to_pt(paragraph_format.space_after) or fallback.space_after,
    )


def _profile_from_cell(cell: object, fallback: StyleProfile) -> StyleProfile:
    for paragraph in cell.paragraphs:
        if paragraph.text.strip():
            return _profile_from_paragraph(paragraph, fallback)
    return fallback


def _apply_style_to_paragraph(paragraph: object, profile: StyleProfile) -> None:
    alignment = name_to_alignment(profile.alignment)
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph_format = paragraph.paragraph_format
    if profile.line_spacing is not None:
        apply_line_spacing(paragraph_format, profile.line_spacing)
    if profile.space_before is not None:
        paragraph_format.space_before = pt_or_none(profile.space_before)
    if profile.space_after is not None:
        paragraph_format.space_after = pt_or_none(profile.space_after)
    if profile.first_line_indent is not None:
        paragraph_format.first_line_indent = pt_or_none(profile.first_line_indent)

    for run in paragraph.runs:
        set_run_font(run, profile.font_name, profile.latin_font_name)
        font_size = font_size_to_pt(profile.font_size_pt)
        if font_size is not None:
            run.font.size = Pt(font_size)
        if profile.bold is not None:
            run.bold = profile.bold


def _reset_table_paragraph_format(paragraph: object) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Pt(0)
    paragraph_format.left_indent = Pt(0)
    paragraph_format.right_indent = Pt(0)


def _preferred_cell_alignment(text: str):
    return WD_ALIGN_PARAGRAPH.CENTER


def _replace_paragraph_text(paragraph: object, text: str) -> None:
    clear_paragraph_runs(paragraph)
    paragraph.add_run(text)


def _normalize_table_caption_text(text: str) -> str:
    normalized = normalize_text(text)
    match = TABLE_CAPTION_PREFIX_RE.match(normalized)
    if match:
        label, number, title = match.groups()
        return f"{label} {number} {title}".rstrip()
    match = ENGLISH_TABLE_CAPTION_PREFIX_RE.match(normalized)
    if match:
        label, number, title = match.groups()
        return f"{label} {number} {title}".rstrip()
    return normalized


def format_table_caption(paragraph: object, style_profile: StyleProfile) -> None:
    normalized = _normalize_table_caption_text(paragraph.text)
    if normalized and normalized != normalize_text(paragraph.text):
        _replace_paragraph_text(paragraph, normalized)
    _apply_style_to_paragraph(paragraph, style_profile)


def is_likely_three_line_table(table: object) -> bool:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders") if tbl_pr is not None else None
    top = border_value(borders, "top") if borders is not None else None
    bottom = border_value(borders, "bottom") if borders is not None else None
    inside_v = border_value(borders, "insideV") if borders is not None else None
    inside_h = border_value(borders, "insideH") if borders is not None else None
    header_has_bottom = False
    first_row_has_top = False
    last_row_has_bottom = False
    if table.rows:
        tr_pr = table.rows[0]._tr.trPr
        tbl_pr_ex = tr_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPrEx") if tr_pr is not None else None
        row_borders = tbl_pr_ex.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders") if tbl_pr_ex is not None else None
        if row_borders is not None and border_value(row_borders, "bottom") not in {None, "nil", "none"}:
            header_has_bottom = True
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.tcPr
            cell_borders = tc_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders") if tc_pr is not None else None
            if cell_borders is not None and border_value(cell_borders, "bottom") not in {None, "nil", "none"}:
                header_has_bottom = True
            if cell_borders is not None and border_value(cell_borders, "top") not in {None, "nil", "none"}:
                first_row_has_top = True
        for cell in table.rows[-1].cells:
            tc_pr = cell._tc.tcPr
            cell_borders = tc_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders") if tc_pr is not None else None
            if cell_borders is not None and border_value(cell_borders, "bottom") not in {None, "nil", "none"}:
                last_row_has_bottom = True
                break
    has_outer_lines = (
        top not in {None, "nil", "none"} and bottom not in {None, "nil", "none"}
    ) or (first_row_has_top and last_row_has_bottom)
    no_vertical_lines = inside_v in {None, "nil", "none"}
    has_table_header_separator = inside_h not in {None, "nil", "none"}
    return has_outer_lines and (header_has_bottom or has_table_header_separator) and no_vertical_lines


def detect_table_style_from_template(template_doc: object) -> TableProfile:
    profile = TableProfile()
    captions = find_table_captions(template_doc)
    if captions:
        profile.caption = _profile_from_paragraph(captions[0], profile.caption)
        profile.used_default = False

    if template_doc.tables:
        table = template_doc.tables[0]
        profile.is_three_line = is_likely_three_line_table(table)
        profile.used_default = False
        if table.rows and table.rows[0].cells:
            profile.header = _profile_from_cell(table.rows[0].cells[0], profile.header)
            profile.header.bold = True if profile.header.bold is None else profile.header.bold
        if len(table.rows) > 1 and table.rows[1].cells:
            profile.body = _profile_from_cell(table.rows[1].cells[0], profile.body)
        elif table.rows and table.rows[0].cells:
            profile.body = _profile_from_cell(table.rows[0].cells[0], profile.body)
    return profile


def apply_three_line_table_style(table: object, style_profile: TableProfile) -> None:
    apply_three_line_borders(
        table,
        top_size=style_profile.borders.top_size,
        header_bottom_size=style_profile.borders.header_bottom_size,
        bottom_size=style_profile.borders.bottom_size,
        color=style_profile.borders.color,
    )


def _clean_cell_text(cell: object) -> str:
    lines = [normalize_text(paragraph.text) for paragraph in cell.paragraphs]
    lines = [line for line in lines if line]
    return "\n".join(lines)


def _extract_table_matrix(table: object) -> list[list[str]]:
    matrix: list[list[str]] = []
    for row in table.rows:
        seen_cells: set[int] = set()
        values: list[str] = []
        for cell in row.cells:
            cell_id = id(cell._tc)
            if cell_id in seen_cells:
                values.append("")
                continue
            seen_cells.add(cell_id)
            values.append(_clean_cell_text(cell))
        matrix.append(values)

    column_count = max((len(row) for row in matrix), default=0)
    if column_count == 0:
        return []
    return [row + [""] * (column_count - len(row)) for row in matrix]


def rebuild_table_from_clean_matrix(document: object, table: object) -> object:
    """Replace a pasted/HTML-styled table with a clean Word table preserving text grid."""

    matrix = _extract_table_matrix(table)
    if not matrix or not matrix[0]:
        return table

    clean_table = document.add_table(rows=len(matrix), cols=len(matrix[0]))
    for row_index, row_values in enumerate(matrix):
        for col_index, text in enumerate(row_values):
            clean_table.cell(row_index, col_index).text = text

    old_tbl = table._tbl
    new_tbl = clean_table._tbl
    old_tbl.addprevious(new_tbl)
    old_tbl.getparent().remove(old_tbl)
    return Table(new_tbl, table._parent)


def _format_table_text(table: object, profile: TableProfile) -> None:
    clear_table_style_and_fills(table)
    for row_index, row in enumerate(table.rows):
        cell_profile = profile.header if row_index == 0 else profile.body
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                _reset_table_paragraph_format(paragraph)
                _apply_style_to_paragraph(paragraph, cell_profile)
                _reset_table_paragraph_format(paragraph)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else _preferred_cell_alignment(paragraph.text)
                for run in paragraph.runs:
                    run.bold = True if row_index == 0 else False


def format_all_tables(document: object, template_profile: object) -> TableFormatResult:
    result = TableFormatResult()
    table_profile: TableProfile = template_profile.table

    captions = find_table_captions(document)
    result.table_captions_detected = len(captions)
    for paragraph in captions:
        format_table_caption(paragraph, table_profile.caption)
        result.table_captions_formatted += 1

    table_note_style = StyleProfile(
        font_name="宋体",
        latin_font_name="Times New Roman",
        font_size_pt=10.5,
        bold=False,
        alignment="LEFT",
        line_spacing=1.0,
        space_before=0,
        space_after=0,
        first_line_indent=0,
    )
    for paragraph in find_table_notes(document):
        _apply_style_to_paragraph(paragraph, table_note_style)

    result.tables_detected = len(document.tables)
    for index, source_table in enumerate(list(document.tables), start=1):
        table = source_table
        if table_has_merged_cells(table):
            result.complex_tables.append(f"第 {index} 个表格包含合并单元格，已按纯文本矩阵尽量重建，建议人工复核。")
        table = rebuild_table_from_clean_matrix(document, table)
        _format_table_text(table, table_profile)
        apply_three_line_table_style(table, table_profile)
        result.tables_formatted += 1
        result.three_line_tables_applied += 1

    if table_profile.used_default:
        result.warnings.append("模板未识别到明确表格样式，系统使用默认中文论文三线表格式。")
    if result.tables_detected:
        result.warnings.append("已按 Excel 清洗思路重建表格：剥离网页/HTML 样式，仅保留纯文本行列结构，再套用论文三线表。")
    return result
